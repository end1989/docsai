"""
Universal file parser for multiple document types.
Supports PDFs, text files, Word docs, emails, eBooks, and more.
"""

import os
import json
import csv
from pathlib import Path
from typing import Dict, List, Tuple, Optional
import mimetypes
from datetime import datetime

# For parsing different file types
try:
    import PyPDF2
    HAS_PDF = True
except ImportError:
    HAS_PDF = False

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    import email
    from email import policy
    from email.parser import BytesParser
    HAS_EMAIL = True
except ImportError:
    HAS_EMAIL = False

try:
    import ebooklib
    from ebooklib import epub
    HAS_EPUB = True
except ImportError:
    HAS_EPUB = False

try:
    from bs4 import BeautifulSoup
    HAS_BS4 = True
except ImportError:
    HAS_BS4 = False

try:
    import markdown
    HAS_MARKDOWN = True
except ImportError:
    HAS_MARKDOWN = False


class FileParser:
    """Universal file parser for various document types."""

    def __init__(self):
        self.parsers = {
            '.txt': self.parse_text,
            '.md': self.parse_markdown,
            '.pdf': self.parse_pdf,
            '.docx': self.parse_docx,
            '.doc': self.parse_docx,  # Older Word files
            '.eml': self.parse_email,
            '.msg': self.parse_email,
            '.epub': self.parse_epub,
            '.html': self.parse_html,
            '.htm': self.parse_html,
            '.json': self.parse_json,
            '.csv': self.parse_csv,
            '.log': self.parse_text,
            '.rtf': self.parse_text,  # Basic RTF as text
        }

    def parse_file(self, filepath: str) -> Optional[Dict[str, str]]:
        """
        Parse a file and return its content and metadata.

        Returns:
            Dict with 'content', 'metadata', and 'error' (if any)
        """
        filepath = Path(filepath)

        if not filepath.exists():
            return {'content': '', 'metadata': {}, 'error': 'File not found'}

        # Get file extension
        ext = filepath.suffix.lower()

        # Get basic metadata
        stat = filepath.stat()
        metadata = {
            'filename': filepath.name,
            'path': str(filepath.absolute()),
            'size_bytes': stat.st_size,
            'modified': datetime.fromtimestamp(stat.st_mtime).isoformat(),
            'extension': ext,
            'mime_type': mimetypes.guess_type(str(filepath))[0] or 'unknown'
        }

        # Choose parser based on extension
        parser = self.parsers.get(ext, self.parse_text)

        try:
            content = parser(filepath)
            return {
                'content': content,
                'metadata': metadata,
                'error': None
            }
        except Exception as e:
            return {
                'content': '',
                'metadata': metadata,
                'error': f"Failed to parse {ext} file: {str(e)}"
            }

    def parse_text(self, filepath: Path) -> str:
        """Parse plain text files."""
        encodings = ['utf-8', 'latin-1', 'cp1252', 'utf-16']

        for encoding in encodings:
            try:
                with open(filepath, 'r', encoding=encoding) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue

        # If all encodings fail, try binary mode and decode with errors='replace'
        with open(filepath, 'rb') as f:
            return f.read().decode('utf-8', errors='replace')

    def parse_markdown(self, filepath: Path) -> str:
        """Parse Markdown files."""
        content = self.parse_text(filepath)

        # Optionally convert to plain text (removing markdown syntax)
        # For now, keep markdown as-is since it's readable
        return content

    def parse_pdf(self, filepath: Path) -> str:
        """Parse PDF files."""
        if not HAS_PDF:
            # Fallback: Just return that PDF parsing is not available
            return f"[PDF: {filepath.name}] - PDF parsing requires PyPDF2 library"

        text_parts = []

        try:
            with open(filepath, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                num_pages = len(pdf_reader.pages)

                text_parts.append(f"[PDF Document: {num_pages} pages]\n")

                for page_num, page in enumerate(pdf_reader.pages, 1):
                    try:
                        text = page.extract_text()
                        if text.strip():
                            text_parts.append(f"\n--- Page {page_num} ---\n{text}")
                    except Exception as e:
                        text_parts.append(f"\n--- Page {page_num} ---\n[Error extracting text: {e}]")

                return "\n".join(text_parts)
        except Exception as e:
            return f"[PDF Error: {e}]"

    def parse_docx(self, filepath: Path) -> str:
        """Parse Word documents."""
        if not HAS_DOCX:
            return f"[DOCX: {filepath.name}] - Word document parsing requires python-docx library"

        try:
            doc = Document(filepath)
            paragraphs = []

            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        paragraphs.append(row_text)

            return '\n\n'.join(paragraphs)
        except Exception as e:
            return f"[DOCX Error: {e}]"

    def parse_email(self, filepath: Path) -> str:
        """Parse email files (.eml, .msg)."""
        if not HAS_EMAIL:
            return f"[Email: {filepath.name}] - Email parsing requires email library"

        try:
            with open(filepath, 'rb') as f:
                msg = BytesParser(policy=policy.default).parse(f)

            parts = []

            # Add headers
            parts.append(f"From: {msg.get('From', 'Unknown')}")
            parts.append(f"To: {msg.get('To', 'Unknown')}")
            parts.append(f"Subject: {msg.get('Subject', 'No Subject')}")
            parts.append(f"Date: {msg.get('Date', 'Unknown')}")
            parts.append("\n---\n")

            # Extract body
            body = ""
            if msg.is_multipart():
                for part in msg.walk():
                    if part.get_content_type() == 'text/plain':
                        body = part.get_payload(decode=True).decode('utf-8', errors='replace')
                        break
                    elif part.get_content_type() == 'text/html' and not body:
                        html_content = part.get_payload(decode=True).decode('utf-8', errors='replace')
                        if HAS_BS4:
                            soup = BeautifulSoup(html_content, 'html.parser')
                            body = soup.get_text()
                        else:
                            body = html_content
            else:
                body = msg.get_payload(decode=True).decode('utf-8', errors='replace')

            parts.append(body)

            return '\n'.join(parts)
        except Exception as e:
            return f"[Email Error: {e}]"

    def parse_epub(self, filepath: Path) -> str:
        """Parse EPUB eBook files."""
        if not HAS_EPUB:
            return f"[EPUB: {filepath.name}] - EPUB parsing requires ebooklib"

        try:
            book = epub.read_epub(filepath)
            text_parts = []

            # Get book metadata
            title = book.get_metadata('DC', 'title')
            author = book.get_metadata('DC', 'creator')

            if title:
                text_parts.append(f"Title: {title[0][0]}")
            if author:
                text_parts.append(f"Author: {author[0][0]}")
            text_parts.append("\n---\n")

            # Extract text from all items
            for item in book.get_items():
                if item.get_type() == ebooklib.ITEM_DOCUMENT:
                    content = item.get_content().decode('utf-8', errors='replace')
                    if HAS_BS4:
                        soup = BeautifulSoup(content, 'html.parser')
                        text = soup.get_text()
                        if text.strip():
                            text_parts.append(text)
                    else:
                        text_parts.append(content)

            return '\n\n'.join(text_parts)
        except Exception as e:
            return f"[EPUB Error: {e}]"

    def parse_html(self, filepath: Path) -> str:
        """Parse HTML files."""
        content = self.parse_text(filepath)

        if HAS_BS4:
            try:
                soup = BeautifulSoup(content, 'html.parser')

                # Remove script and style elements
                for script in soup(["script", "style"]):
                    script.decompose()

                # Get text
                text = soup.get_text()

                # Break into lines and remove leading/trailing space
                lines = (line.strip() for line in text.splitlines())
                # Break multi-headlines into a line each
                chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
                # Drop blank lines
                text = '\n'.join(chunk for chunk in chunks if chunk)

                return text
            except Exception:
                return content
        else:
            return content

    def parse_json(self, filepath: Path) -> str:
        """Parse JSON files into readable text."""
        try:
            with open(filepath, 'r', encoding='utf-8') as f:
                data = json.load(f)

            # Pretty print the JSON
            return json.dumps(data, indent=2, ensure_ascii=False)
        except Exception as e:
            # If JSON parsing fails, treat as text
            return self.parse_text(filepath)

    def parse_csv(self, filepath: Path) -> str:
        """Parse CSV files into readable text."""
        try:
            rows = []
            with open(filepath, 'r', encoding='utf-8') as f:
                reader = csv.reader(f)
                for row in reader:
                    rows.append(' | '.join(row))

            return '\n'.join(rows)
        except Exception:
            # If CSV parsing fails, treat as text
            return self.parse_text(filepath)


def scan_directory(directory: str, file_types: List[str] = None,
                  recursive: bool = True) -> List[str]:
    """
    Scan a directory for files to process.

    Args:
        directory: Path to directory to scan
        file_types: List of extensions to include (e.g., ['pdf', 'txt'])
                   If None or includes 'all', all supported types are included
        recursive: Whether to scan subdirectories

    Returns:
        List of file paths
    """
    directory = Path(directory)

    if not directory.exists() or not directory.is_dir():
        return []

    parser = FileParser()
    supported_extensions = set(parser.parsers.keys())

    # Determine which extensions to look for
    if file_types is None or 'all' in file_types:
        target_extensions = supported_extensions
    else:
        # Add dots to extensions if not present
        target_extensions = {f'.{ft}' if not ft.startswith('.') else ft
                           for ft in file_types}
        # Filter to only supported extensions
        target_extensions = target_extensions & supported_extensions

    files = []

    if recursive:
        for ext in target_extensions:
            pattern = f"**/*{ext}"
            files.extend(str(p) for p in directory.glob(pattern) if p.is_file())
    else:
        for ext in target_extensions:
            pattern = f"*{ext}"
            files.extend(str(p) for p in directory.glob(pattern) if p.is_file())

    return sorted(set(files))  # Remove duplicates and sort


def parse_directory(directory: str, file_types: List[str] = None,
                   recursive: bool = True) -> List[Dict]:
    """
    Parse all supported files in a directory.

    Returns:
        List of parsed documents with content and metadata
    """
    files = scan_directory(directory, file_types, recursive)
    parser = FileParser()

    results = []
    for filepath in files:
        result = parser.parse_file(filepath)
        if result and result.get('content'):
            results.append(result)

    return results